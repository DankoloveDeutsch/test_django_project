# module_app/utils/document_generator.py
"""
Генерация документов (PDF, DOCX, HTML)
"""

import os
import re
import io
from datetime import datetime
from django.conf import settings
from django.template import Template, Context
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
import pdfkit
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def render_document_template(content, context):
    """
    Рендеринг шаблона документа с подстановкой переменных

    Args:
        content: текст шаблона с переменными {{ variable }}
        context: словарь с значениями переменных

    Returns:
        str: отрендеренный текст
    """
    # Регулярное выражение для поиска переменных {{ variable }}
    pattern = r'\{\{\s*(\w+)\s*\}\}'

    def replace_variable(match):
        var_name = match.group(1)
        return str(context.get(var_name, match.group(0)))

    return re.sub(pattern, replace_variable, content)


def generate_pdf_from_html(html_content, filename=None):
    """
    Генерация PDF из HTML

    Args:
        html_content: HTML контент
        filename: имя файла (опционально)

    Returns:
        bytes: PDF контент
    """
    # Настройки для wkhtmltopdf
    options = {
        'page-size': 'A4',
        'margin-top': '1.5cm',
        'margin-right': '1.5cm',
        'margin-bottom': '1.5cm',
        'margin-left': '1.5cm',
        'encoding': 'UTF-8',
        'no-outline': None,
        'enable-local-file-access': None
    }

    # Добавляем базовые стили
    styled_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            body {{
                font-family: 'DejaVu Sans', Arial, sans-serif;
                font-size: 12pt;
                line-height: 1.4;
                color: #333;
            }}
            h1 {{
                font-size: 18pt;
                text-align: center;
                margin-bottom: 20px;
            }}
            h2 {{
                font-size: 16pt;
                margin-top: 20px;
                margin-bottom: 10px;
            }}
            table {{
                border-collapse: collapse;
                width: 100%;
                margin: 15px 0;
            }}
            th, td {{
                border: 1px solid #ddd;
                padding: 8px;
                text-align: left;
            }}
            th {{
                background-color: #f5f5f5;
            }}
            .signature {{
                margin-top: 50px;
                display: flex;
                justify-content: space-between;
            }}
            .footer {{
                position: fixed;
                bottom: 0;
                width: 100%;
                text-align: center;
                font-size: 9pt;
                color: #999;
                padding: 10px 0;
            }}
        </style>
    </head>
    <body>
        {html_content}
        <div class="footer">
            Сгенерировано в системе "Управление табельным учетом" {datetime.now().strftime('%d.%m.%Y %H:%M')}
        </div>
    </body>
    </html>
    """

    try:
        # Используем pdfkit с wkhtmltopdf
        pdf_content = pdfkit.from_string(styled_html, False, options=options)
        return pdf_content
    except Exception as e:
        # Если pdfkit недоступен, используем reportlab
        return generate_pdf_reportlab(html_content)


def generate_pdf_reportlab(html_content):
    """
    Генерация PDF с помощью ReportLab (альтернативный метод)

    Args:
        html_content: HTML контент

    Returns:
        bytes: PDF контент
    """
    buffer = io.BytesIO()

    # Создаем PDF документ
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            rightMargin=2 * cm, leftMargin=2 * cm,
                            topMargin=2 * cm, bottomMargin=2 * cm)

    # Стили
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='Center', alignment=TA_CENTER, fontSize=14, leading=18))
    styles.add(ParagraphStyle(name='Right', alignment=TA_RIGHT))

    # Удаляем HTML теги для ReportLab
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')
    text = soup.get_text()

    # Создаем элементы
    story = []

    # Добавляем текст
    for line in text.split('\n'):
        if line.strip():
            story.append(Paragraph(line, styles['Normal']))
            story.append(Spacer(1, 0.2 * cm))

    # Строим PDF
    doc.build(story)

    buffer.seek(0)
    return buffer.getvalue()


def generate_docx(html_content, filename=None):
    """
    Генерация DOCX документа

    Args:
        html_content: HTML контент
        filename: имя файла (опционально)

    Returns:
        bytes: DOCX контент
    """
    from bs4 import BeautifulSoup

    document = Document()

    # Устанавливаем поля
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Парсим HTML
    soup = BeautifulSoup(html_content, 'html.parser')

    # Добавляем содержимое
    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'table', 'ul', 'ol']):
        if element.name == 'h1':
            p = document.add_heading(element.get_text(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif element.name == 'h2':
            document.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            document.add_heading(element.get_text(), level=3)
        elif element.name == 'p':
            p = document.add_paragraph(element.get_text())
            # Проверяем выравнивание
            if element.get('style'):
                if 'text-align: center' in element['style']:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif 'text-align: right' in element['style']:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif element.name == 'table':
            # Создаем таблицу
            rows = element.find_all('tr')
            if rows:
                table = document.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                table.style = 'Table Grid'

                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.get_text().strip()
                        if cell.name == 'th':
                            # Жирный шрифт для заголовков
                            run = table.cell(i, j).paragraphs[0].runs[0]
                            run.bold = True
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li'):
                document.add_paragraph(li.get_text(), style='List Bullet')

    # Сохраняем в буфер
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()


def generate_document_pdf(content, document_type, context):
    """
    Генерация PDF документа с использованием шаблона

    Args:
        content: шаблон документа
        document_type: тип документа
        context: контекст для подстановки

    Returns:
        bytes: PDF контент
    """
    # Рендерим шаблон
    rendered = render_document_template(content, context)

    # Добавляем заголовок в зависимости от типа
    titles = {
        'employment_order': 'ПРИКАЗ О ПРИЕМЕ НА РАБОТУ',
        'dismissal_order': 'ПРИКАЗ ОБ УВОЛЬНЕНИИ',
        'vacation_order': 'ПРИКАЗ О ПРЕДОСТАВЛЕНИИ ОТПУСКА',
        'employment_contract': 'ТРУДОВОЙ ДОГОВОР',
        'vacation_notice': 'УВЕДОМЛЕНИЕ ОБ ОТПУСКЕ',
        'certificate': 'СПРАВКА'
    }

    title = titles.get(document_type, 'ДОКУМЕНТ')

    # Формируем HTML
    html_content = f"""
    <h1>{title}</h1>
    <div class="content">
        {rendered}
    </div>
    <div class="signature">
        <div>Руководитель: __________________</div>
        <div>Сотрудник: __________________</div>
    </div>
    """

    return generate_pdf_from_html(html_content)


def generate_employee_card(employee):
    """
    Генерация карточки сотрудника в PDF

    Args:
        employee: объект EmployeeProfile

    Returns:
        bytes: PDF контент
    """
    html_content = f"""
    <h1>ЛИЧНАЯ КАРТОЧКА СОТРУДНИКА</h1>

    <h2>Общая информация</h2>
    <table>
        <tr><th>Табельный номер:</th><td>{employee.employee_code or '—'}</td></tr>
        <tr><th>ФИО:</th><td>{employee.full_name}</td></tr>
        <tr><th>Должность:</th><td>{employee.position or '—'}</td></tr>
        <tr><th>Отдел:</th><td>{employee.department or '—'}</td></tr>
        <tr><th>Дата рождения:</th><td>{employee.birth_date.strftime('%d.%m.%Y') if employee.birth_date else '—'}</td></tr>
        <tr><th>Дата приема:</th><td>{employee.hire_date.strftime('%d.%m.%Y') if employee.hire_date else '—'}</td></tr>
        <tr><th>Телефон:</th><td>{employee.phone or '—'}</td></tr>
        <tr><th>Email:</th><td>{employee.user.email}</td></tr>
    </table>

    <h2>Документы</h2>
    <table>
        <tr><th>Тип</th><th>Название</th><th>Срок действия</th></tr>
        {''.join(f'<tr><td>{doc.get_document_type_display()}</td><td>{doc.title}</td><td>{doc.expiry_date.strftime("%d.%m.%Y") if doc.expiry_date else "—"}</td></tr>' for doc in employee.personal_documents.all()[:10])}
    </table>

    <h2>Статистика</h2>
    <table>
        <tr><th>Всего часов:</th><td>{employee.total_hours}</td></tr>
        <tr><th>Документов:</th><td>{employee.personal_documents.count()}</td></tr>
        <tr><th>Напоминаний:</th><td>{employee.reminders.filter(is_completed=False).count()}</td></tr>
    </table>
    """

    return generate_pdf_from_html(html_content)